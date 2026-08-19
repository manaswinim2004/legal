from pathlib import Path

import pymupdf
import pytesseract
from PIL import Image
from docx import Document


# Windows Tesseract
pytesseract.pytesseract.tesseract_cmd = (
    r"C:\Program Files\Tesseract-OCR\tesseract.exe"
)


class DocumentProcessor:
    """
    Generic document processor.

    Supported:
        PDF
        DOC
        DOCX
        TXT
        PNG
        JPG
        JPEG

    Output:
        {
            "source": "...",
            "file_type": "...",
            "text": "...",
            "blocks": [...]
        }
    """

    SUPPORTED_EXTENSIONS = {
        ".pdf",
        ".doc",
        ".docx",
        ".txt",
        ".png",
        ".jpg",
        ".jpeg",
    }

    def process(self, file_path: str) -> dict:
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(
                f"File not found: {file_path}"
            )

        extension = path.suffix.lower()

        if extension not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"Unsupported file type: {extension}"
            )

        if extension == ".pdf":
            blocks = self._process_pdf(path)

        elif extension == ".docx":
            blocks = self._process_docx(path)

        elif extension == ".doc":
            blocks = self._process_doc(path)

        elif extension == ".txt":
            blocks = self._process_txt(path)

        else:
            blocks = self._process_image(path)

        text = "\n".join(
            block["text"]
            for block in blocks
            if block["text"].strip()
        )

        return {
            "source": path.name,
            "file_type": extension,
            "text": text.strip(),
            "blocks": blocks,
        }

    # ---------------------------------------------------------
    # PDF
    # ---------------------------------------------------------

    def _process_pdf(self, path: Path) -> list[dict]:
        doc = pymupdf.open(path)

        blocks = []

        for page_number, page in enumerate(doc, start=1):

            text = page.get_text("text").strip()

            if text:
                page_lines = text.splitlines()

                for line in page_lines:
                    line = self._clean_text(line)

                    if line:
                        blocks.append({
                            "type": "paragraph",
                            "text": line,
                            "page": page_number,
                            "method": "text_extraction",
                        })

                continue

            print(
                f"No text found on PDF page {page_number}. "
                f"Running OCR..."
            )

            pix = page.get_pixmap(
                matrix=pymupdf.Matrix(2, 2)
            )

            image = Image.frombytes(
                "RGB",
                [pix.width, pix.height],
                pix.samples,
            )

            ocr_text = pytesseract.image_to_string(
                image
            ).strip()

            for line in ocr_text.splitlines():
                line = self._clean_text(line)

                if line:
                    blocks.append({
                        "type": "paragraph",
                        "text": line,
                        "page": page_number,
                        "method": "ocr",
                    })

        doc.close()

        return blocks

    # ---------------------------------------------------------
    # DOCX
    # ---------------------------------------------------------

    def _process_docx(self, path: Path) -> list[dict]:
        document = Document(path)

        blocks = []

        for paragraph in document.paragraphs:

            text = self._clean_text(paragraph.text)

            if not text:
                continue

            style = paragraph.style.name.lower()

            if "heading" in style:
                block_type = "heading"
            elif "title" in style:
                block_type = "title"
            else:
                block_type = "paragraph"

            blocks.append({
                "type": block_type,
                "text": text,
                "page": None,
                "method": "docx",
                "style": paragraph.style.name,
            })

        # Tables
        for table in document.tables:

            for row in table.rows:

                cells = []

                for cell in row.cells:
                    cell_text = self._clean_text(
                        cell.text
                    )

                    if cell_text:
                        cells.append(cell_text)

                if cells:
                    blocks.append({
                        "type": "table_row",
                        "text": " | ".join(cells),
                        "page": None,
                        "method": "docx",
                    })

        return blocks

    # ---------------------------------------------------------
    # Legacy DOC
    # ---------------------------------------------------------

    def _process_doc(self, path: Path) -> list[dict]:
        """
        Legacy .doc files require LibreOffice conversion.
        """

        import shutil
        import subprocess
        import tempfile

        libreoffice = shutil.which("soffice")

        if not libreoffice:
            raise RuntimeError(
                "DOC files require LibreOffice. "
                "Install LibreOffice and ensure "
                "'soffice' is available in PATH."
            )

        with tempfile.TemporaryDirectory() as temp_dir:

            subprocess.run(
                [
                    libreoffice,
                    "--headless",
                    "--convert-to",
                    "docx",
                    "--outdir",
                    temp_dir,
                    str(path),
                ],
                check=True,
                capture_output=True,
                text=True,
            )

            converted = (
                Path(temp_dir)
                / f"{path.stem}.docx"
            )

            if not converted.exists():
                raise RuntimeError(
                    "DOC → DOCX conversion failed."
                )

            return self._process_docx(converted)

    # ---------------------------------------------------------
    # TXT
    # ---------------------------------------------------------

    def _process_txt(self, path: Path) -> list[dict]:

        text = path.read_text(
            encoding="utf-8",
            errors="replace",
        )

        blocks = []

        for line in text.splitlines():

            line = self._clean_text(line)

            if line:
                blocks.append({
                    "type": "paragraph",
                    "text": line,
                    "page": None,
                    "method": "txt",
                })

        return blocks

    # ---------------------------------------------------------
    # IMAGE
    # ---------------------------------------------------------

    def _process_image(self, path: Path) -> list[dict]:

        image = Image.open(path)

        text = pytesseract.image_to_string(
            image
        ).strip()

        blocks = []

        for line in text.splitlines():

            line = self._clean_text(line)

            if line:
                blocks.append({
                    "type": "paragraph",
                    "text": line,
                    "page": 1,
                    "method": "ocr",
                })

        return blocks

    # ---------------------------------------------------------
    # Utility
    # ---------------------------------------------------------

    @staticmethod
    def _clean_text(text: str) -> str:
        return " ".join(text.split()).strip()